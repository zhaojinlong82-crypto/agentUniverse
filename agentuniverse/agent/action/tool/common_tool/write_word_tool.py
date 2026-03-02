import os
import json
from typing import Any, Dict

from agentuniverse.agent.action.tool.tool import Tool


class WriteWordDocumentTool(Tool):
    def execute(self, file_path: str, content: str = "", append: bool = False) -> str:
        directory = os.path.dirname(file_path)
        if directory and not os.path.exists(directory):
            try:
                os.makedirs(directory, exist_ok=True)
            except Exception as e:
                return json.dumps(
                    {"error": f"Failed to create directory: {str(e)}", "file_path": file_path, "status": "error"}
                )

        try:
            from docx import Document  # type: ignore
        except ImportError as e:
            return json.dumps(
                {
                    "error": f"python-docx is required to write Word documents: {str(e)}",
                    "file_path": file_path,
                    "status": "error",
                }
            )

        if not file_path.lower().endswith(".docx"):
            return json.dumps(
                {"error": "The target file must have a .docx extension.", "file_path": file_path, "status": "error"}
            )

        document = None
        if append and os.path.exists(file_path):
            try:
                document = Document(file_path)
            except Exception as e:
                return json.dumps(
                    {"error": f"Failed to load existing document: {str(e)}", "file_path": file_path, "status": "error"}
                )
        else:
            document = Document()

        try:
            document.add_paragraph(content)
            document.save(file_path)
            file_size = os.path.getsize(file_path)
            return json.dumps(
                {
                    "file_path": file_path,
                    "bytes_written": len(content.encode("utf-8")),
                    "file_size": file_size,
                    "append_mode": append,
                    "status": "success",
                }
            )
        except Exception as e:
            return json.dumps(
                {"error": f"Failed to write document: {str(e)}", "file_path": file_path, "status": "error"}
            )
